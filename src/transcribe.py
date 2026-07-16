"""Gemini API への文字起こしリクエストと、Word ファイル生成"""
from __future__ import annotations

import re
import time
from pathlib import Path

from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


# ======================================================================
# プロンプト部品
#
# v1.3.0: プロンプトを固定文字列から「部品の組み立て」に変更した。
#   - 整形ルール(従来) / 逐語ルール(新設) を切り替え可能
#   - 参加者名簿(roster)があれば話者ラベルを実名にする
# ======================================================================

_RULES_CLEANUP = """- 内容は正確に一言一句書き起こす(情報の欠落・改変なし)
- フィラー(「えー」「あのー」「えっと」「まあ」「そのー」など)は適宜削除
- 言い淀み・不要な繰り返しは整理し、読みやすい日本語に整える
- 話者の意図・固有名詞・数字は正確に保つ
- 聞き取れなかった箇所は [不明] と記す"""

_RULES_VERBATIM = """【逐語(一言一句)ルール・最重要】
- 発言は一切要約・整文しない。話されたとおりに書き起こす
- 「えー」「あのー」「えっと」等の言いよどみ(フィラー)、言い直し、繰り返しも省略せずそのまま残す
- 聞き取れない箇所は推測で補完せず、必ず「(聴取不能)」と表記する
- 音声に存在しない語句を付け加えない。文法的に不自然でも直さない
- 固有名詞が不明瞭な場合は、聞こえたとおりに表記する(勝手に「正しい」名前に直さない)
- 実際に発話された内容だけを書く。音声上で繰り返されていない限り、出力で同じ言葉を繰り返さない"""

_RULES_TS = """- **段落の先頭に必ず [MM:SS] 形式のタイムスタンプを入れる**
  (音声内のその段落が始まる時刻、ゼロ埋め2桁、例: [00:00], [03:45])
- 段落は話題のまとまり・明確な区切り・または30秒~2分ごとに分ける"""

_RULES_DIAR = """- **話者が変わるごとに新しい行として書き出す**
- **各行の冒頭に必ず以下の形式を入れる**:
  [MM:SS] 【話者ラベル】 本文...
  (時刻はゼロ埋め2桁の分:秒。ミリ秒や1/100秒は付けない)
- 短い相づち(「はい」「うん」程度)は前の発言と同じ行に含めてよい"""

_TAIL = "- 説明や前置き、Markdown 装飾は不要。書き起こし本文のみを出力する。"


def _roster_block(roster: str) -> str:
    return f"""この音声の参加者は以下のとおり事前に判明しています。話者の判別には、声質に加えて、
発言内容(名乗り・指名・役職に応じた発言内容)も手がかりにしてください。

【参加者名簿】
{roster.strip()}

- 話者ラベルは名簿にある呼称を【】で囲んで用いる(例: 【議長(理事長)】【佐藤理事】)
- どうしても特定できない場合のみ【発言者不明】、複数人の発話が重なる場合は【発言者複数・重複】とする。
  無理に名簿の誰かに割り当てるより、不明と明示するほうが望ましい"""


_LABEL_ABC = """- 話者ラベル(発言者A, 発言者B, 発言者C...)は声の特徴で識別し、同じ人物には常に同じラベルを使う
- 話者を特定できない場合は【発言者不明】とする"""

_EXAMPLE_DIAR = """出力例:
[00:00] 【発言者A】 本日はお忙しい中お集まりいただきありがとうございます。それでは議事を始めます。
[00:25] 【発言者B】 すみません、確認ですが、前回の議事録は配布済みですか?
[00:32] 【発言者A】 はい、配布済みです。修正点も反映されています。"""

_EXAMPLE_TS = """出力例:
[00:00] 本日の会議を開始します。まず議題ですが、予算と人事の二点になります。
[00:42] それでは一つ目、来期予算について議論を始めます。
[03:15] 続いて人事についての検討に移ります。"""


def build_prompt(
    with_timestamps: bool,
    with_diarization: bool,
    roster: str = "",
    verbatim: bool = False,
) -> str:
    """設定に応じて文字起こしプロンプトを組み立てる。"""
    parts: list[str] = ["この音声を日本語で書き起こしてください。", ""]

    if with_diarization and roster.strip():
        parts += [_roster_block(roster), ""]

    parts.append("ルール:")
    if with_diarization:
        parts.append(_RULES_DIAR)
        if not roster.strip():
            parts.append(_LABEL_ABC)
    elif with_timestamps:
        parts.append(_RULES_TS)

    parts.append(_RULES_VERBATIM if verbatim else _RULES_CLEANUP)
    if verbatim:
        # 逐語モードでも段落・句読点の最低限の整形は許可する
        parts.append("- 句読点は聞こえたとおりの区切りで付けてよい(内容の変更は不可)")
    parts.append(_TAIL)

    if with_diarization and not roster.strip():
        parts += ["", _EXAMPLE_DIAR]
    elif with_timestamps and not with_diarization:
        parts += ["", _EXAMPLE_TS]

    return "\n".join(parts)


# 後方互換(旧コードが参照していた固定プロンプト)
PROMPT_PLAIN = build_prompt(False, False)
PROMPT_TIMESTAMP = build_prompt(True, False)
PROMPT_DIARIZATION = build_prompt(True, True)


DIARIZATION_NOTE = (
    "※ 本文中の話者ラベル(発言者A/B/C 等)は分割チャンクごとに声で識別しているため、"
    "チャンク境界をまたぐと同一人物が別ラベルになる可能性があります。"
    "長い音声では発言者の対応関係を読みながらご確認ください。"
)

ROSTER_NOTE = (
    "※ 本文中の話者ラベルは、入力された参加者名簿に基づきAIが推定したものです。"
    "特定できなかった発言は【発言者不明】等と表記されます。"
    "最終利用の前に、必ず原音声と突き合わせてご確認ください。"
)


# ======================================================================
# 暴走ループ(縮退出力)の検出
#
# v1.3.0: temperature 0.0 相当の生成で、同じパターンを無限に繰り返す
# 「暴走ループ」が発生する事例が確認された(検証実験 2026-07)。
# 検出したら temperature を段階的に上げて再生成する。
# ======================================================================

_LOOP_THRESHOLD = 15
_INLINE_LOOP = re.compile(r"(.{1,12}?)\1{%d,}" % (_LOOP_THRESHOLD - 1))

# 暴走時に段階的に試す temperature
_TEMPERATURES = (0.0, 0.3, 0.7)


def is_degenerate(text: str) -> bool:
    """暴走ループの検出(3段構え)

    (a) 行内の短いパターンの大量繰り返し(例: 「お、お、お、…」が1行に続く)
    (b) 行の種類が極端に少ない(例: 2行が交互に数千回続く)
    (c) 同一行の連続
    """
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        return True  # 空出力も失敗扱い

    for ln in lines:
        if len(ln) >= _LOOP_THRESHOLD * 2 and _INLINE_LOOP.search(ln):
            return True

    if len(lines) >= 40 and len(set(lines)) / len(lines) < 0.2:
        return True

    run = 1
    for prev, cur in zip(lines, lines[1:]):
        if cur == prev:
            run += 1
            if run >= _LOOP_THRESHOLD:
                return True
        else:
            run = 1
    return False


# ======================================================================
# タイムスタンプ処理(v1.2 から変更なし)
# ======================================================================

# Gemini が返す [MM:SS] または [M:SS](チャンク内相対時刻)を捕捉
_TS_REL_PATTERN = re.compile(r"\[(\d{1,2}):(\d{2})\]")

# 絶対時刻 + 任意の話者ラベルを段落冒頭で検出(太字化用)
_TS_LEAD_PATTERN = re.compile(
    r"^(\[\d{1,2}:\d{2}(?::\d{2})?\])"   # group(1): 時刻
    r"\s*"
    r"(【[^】]+】)?"                       # group(2): 話者ラベル(任意)
    r"\s*"
    r"(.*)"                                # group(3): 本文
)


def shift_timestamps(text: str, offset_seconds: int) -> str:
    """テキスト中の [MM:SS] (チャンク内相対時刻) を offset_seconds 加算して
    [HH:MM:SS] (元音声内の絶対時刻) に変換する。
    """
    def repl(m: "re.Match[str]") -> str:
        total = int(m.group(1)) * 60 + int(m.group(2)) + offset_seconds
        h = total // 3600
        mm = (total % 3600) // 60
        ss = total % 60
        return f"[{h:02d}:{mm:02d}:{ss:02d}]"
    return _TS_REL_PATTERN.sub(repl, text)


# ======================================================================
# 文字起こし本体
# ======================================================================

def transcribe_audio(
    client: genai.Client,
    audio_path: Path,
    model: str,
    with_timestamps: bool = False,
    with_diarization: bool = False,
    roster: str = "",
    verbatim: bool = False,
    max_retries: int = 3,
    on_log=None,
) -> str:
    """1チャンクをGeminiで文字起こし。

    - 通信エラー等は指数バックオフで再試行(従来どおり)
    - 暴走ループを検出したら temperature を上げて再生成(v1.3.0)
    """
    last_error: Exception | None = None
    prompt = build_prompt(with_timestamps, with_diarization, roster, verbatim)

    def log(msg: str) -> None:
        if on_log:
            on_log(msg)

    for attempt in range(max_retries):
        try:
            uploaded = client.files.upload(file=str(audio_path))

            waited = 0
            while uploaded.state.name == "PROCESSING":
                time.sleep(1)
                waited += 1
                if waited > 300:  # 5分タイムアウト
                    raise RuntimeError("ファイル処理がタイムアウトしました。")
                uploaded = client.files.get(name=uploaded.name)

            if uploaded.state.name == "FAILED":
                raise RuntimeError(f"アップロード失敗: {audio_path.name}")

            text = ""
            for temp in _TEMPERATURES:
                response = client.models.generate_content(
                    model=model,
                    contents=[uploaded, prompt],
                    config=types.GenerateContentConfig(temperature=temp),
                )
                text = (response.text or "").strip()
                if not is_degenerate(text):
                    break
                log(f"  ※ 出力の暴走を検出 (temp={temp})。設定を変えて再生成します...")
                time.sleep(2)
            else:
                # 全温度で暴走。最後の結果に警告を付けて返す(処理は止めない)
                log("  ※ 再生成でも暴走が解消しませんでした。このチャンクは要確認です。")
                text = f"【警告: このチャンクの出力は不安定でした。原音声を確認してください】\n{text}"

            try:
                client.files.delete(name=uploaded.name)
            except Exception:
                pass

            return text

        except Exception as e:
            last_error = e
            if attempt < max_retries - 1:
                wait = 2 ** attempt
                time.sleep(wait)

    assert last_error is not None
    raise last_error


# ======================================================================
# Word 出力
# ======================================================================

def _ensure_japanese_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "游明朝"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "游明朝")


def write_docx(
    transcripts: list[str],
    output_path: Path,
    title: str,
    note: str | None = None,
) -> None:
    """チャンクごとの文字起こし文字列リストを1つの .docx にまとめる。

    段落冒頭が [HH:MM:SS] や 【発言者X】 の形式なら、その部分を太字にする。
    note が指定されていれば、本文先頭に注意書き(斜体・灰色)を入れる。
    """
    doc = Document()
    _ensure_japanese_font(doc)
    doc.add_heading(title, level=1)

    if note:
        p = doc.add_paragraph()
        run = p.add_run(note)
        run.italic = True
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
        doc.add_paragraph()  # 空行で区切る

    for chunk_text in transcripts:
        for para in chunk_text.split("\n"):
            para = para.strip()
            if not para:
                continue
            m = _TS_LEAD_PATTERN.match(para)
            if m:
                ts, speaker, body = m.group(1), m.group(2), m.group(3)
                p = doc.add_paragraph()
                ts_run = p.add_run(ts + " ")
                ts_run.bold = True
                if speaker:
                    sp_run = p.add_run(speaker + " ")
                    sp_run.bold = True
                if body:
                    p.add_run(body)
            else:
                doc.add_paragraph(para)

    doc.save(output_path)
